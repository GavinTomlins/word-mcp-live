"""Build a complete .docx from a Markdown subset in one pass (ADR 0004/0006).

Supported subset: ``#``–``######`` headings, paragraphs, ``**bold**``,
``*italic*``, ``***bold italic***``, `` `code` ``, ``[text](url)`` links,
nested ``-``/``*``/``+`` bullet lists and ``1.`` numbered lists (2-space
indent per level), GitHub-style pipe tables (header row bold), and ``---``
horizontal rules rendered as a bottom-border paragraph (never a table used
as a divider).

With ``template=`` the document is built on a copy of an existing .docx:
its styles, headers/footers, and page setup are inherited, its body content
is cleared, and the Markdown content is written using the template's own
style definitions — with fallbacks for style names the template does not
define (lists degrade to indented bullet/number text, tables to manually
bordered grids).
"""

import re
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT

_INLINE_RE = re.compile(
    r"(\*\*\*(?P<bi>.+?)\*\*\*"
    r"|\*\*(?P<b>.+?)\*\*"
    r"|\*(?P<i>[^*]+?)\*"
    r"|`(?P<code>[^`]+?)`"
    r"|\[(?P<label>[^\]]+)\]\((?P<url>[^)\s]+)\))"
)

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_NUMBER_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|(\s*:?-{3,}:?\s*\|)+\s*$")
_HR_RE = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph, text: str, bold: bool = False) -> None:
    """Append markdown inline content to a paragraph as formatted runs."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = bold or None
        if m.group("bi") is not None:
            run = paragraph.add_run(m.group("bi"))
            run.bold, run.italic = True, True
        elif m.group("b") is not None:
            run = paragraph.add_run(m.group("b"))
            run.bold = True
        elif m.group("i") is not None:
            run = paragraph.add_run(m.group("i"))
            run.italic = True
            run.bold = bold or None
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Courier New"
            run.bold = bold or None
        elif m.group("label") is not None:
            _add_hyperlink(paragraph, m.group("label"), m.group("url"))
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold or None


def _add_horizontal_rule(doc) -> None:
    paragraph = doc.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pbdr.append(bottom)
    ppr.append(pbdr)


def _list_style(kind: str, level: int) -> str:
    base = "List Bullet" if kind == "bullet" else "List Number"
    return base if level == 0 else f"{base} {min(level, 2) + 1}"


def _list_paragraph(doc, kind: str, level: int):
    """Add a list paragraph, falling back when the template lacks list styles.

    Returns ``(paragraph, styled)``; when ``styled`` is False the caller
    must render the marker as literal text.
    """
    base = "List Bullet" if kind == "bullet" else "List Number"
    for name in (_list_style(kind, level), base, "List Paragraph"):
        try:
            return doc.add_paragraph(style=name), True
        except KeyError:
            continue
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return paragraph, False


def _apply_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _add_table(doc, rows: list[list[str]], has_header: bool) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        # Template has no Table Grid style — draw the grid directly.
        _apply_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            paragraph = table.rows[r_idx].cells[c_idx].paragraphs[0]
            _add_inline(paragraph, cell_text, bold=has_header and r_idx == 0)


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [c.strip().replace("\\|", "|") for c in re.split(r"(?<!\\)\|", stripped)]


def _clear_body(doc) -> None:
    """Remove all body content while preserving the trailing sectPr
    (page setup, headers/footers)."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def markdown_to_document(
    markdown: str,
    filename: str,
    title: str | None = None,
    author: str | None = None,
    template: str | None = None,
) -> dict:
    """Build ``filename`` from Markdown; returns creation statistics.

    With ``template``, the output starts as a copy of that .docx: styles,
    headers/footers and page setup are inherited and the template's body
    content is cleared before the Markdown content is written.
    """
    if template:
        shutil.copyfile(template, filename)
        doc = Document(filename)
        _clear_body(doc)
    else:
        doc = Document()
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    stats = {"headings": 0, "paragraphs": 0, "list_items": 0, "tables": 0, "rules": 0}
    lines = markdown.splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            _add_inline(doc.add_paragraph(), " ".join(paragraph_buffer))
            stats["paragraphs"] += 1
            paragraph_buffer.clear()

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            flush_paragraph()
            level = len(m.group(1))
            try:
                heading = doc.add_heading("", level=level)
                _add_inline(heading, m.group(2).strip())
            except KeyError:
                # Template lacks this heading style — bold paragraph instead.
                heading = doc.add_paragraph()
                _add_inline(heading, m.group(2).strip(), bold=True)
            stats["headings"] += 1
            i += 1
            continue

        if _HR_RE.match(line) and not paragraph_buffer:
            _add_horizontal_rule(doc)
            stats["rules"] += 1
            i += 1
            continue

        if _TABLE_ROW_RE.match(line):
            flush_paragraph()
            table_rows = []
            has_header = False
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                if _TABLE_SEP_RE.match(lines[i]):
                    has_header = len(table_rows) == 1
                else:
                    table_rows.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, table_rows, has_header)
            stats["tables"] += 1
            continue

        bullet_m = _BULLET_RE.match(line)
        number_m = None if bullet_m else _NUMBER_RE.match(line)
        if bullet_m or number_m:
            flush_paragraph()
            kind = "bullet" if bullet_m else "number"
            m = bullet_m or number_m
            level = len(m.group(1)) // 2
            item_text = m.group(2 if bullet_m else 3).strip()
            paragraph, styled = _list_paragraph(doc, kind, level)
            if not styled:
                marker = "• " if bullet_m else f"{number_m.group(2)}. "
                paragraph.add_run(marker)
            _add_inline(paragraph, item_text)
            stats["list_items"] += 1
            i += 1
            continue

        paragraph_buffer.append(line.strip())
        i += 1

    flush_paragraph()
    doc.save(filename)
    return stats
