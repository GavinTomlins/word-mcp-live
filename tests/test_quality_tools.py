"""Tests for validation, markdown round-trip and batch creation tools."""

import asyncio
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

from word_mcp_live_gavintomlins.core.markdown_write import markdown_to_document
from word_mcp_live_gavintomlins.core.markdown_read import document_to_markdown
from word_mcp_live_gavintomlins.core.validation import validate_docx
from word_mcp_live_gavintomlins.tools import quality_tools

SAMPLE_MARKDOWN = """\
# Quarterly Report

## Summary

This report covers **revenue** and *growth* for `Q2` with a
[reference](https://example.com/data).

- First finding
- Second finding
  - Nested detail
1. Step one
2. Step two

| Region | Revenue |
| --- | --- |
| North | 100 |
| South | 200 |

---

Closing paragraph.
"""


def test_markdown_round_trip(tmp_path: Path):
    path = str(tmp_path / "report.docx")
    stats = markdown_to_document(SAMPLE_MARKDOWN, path, title="Report", author="Tester")
    assert stats["headings"] == 2
    assert stats["tables"] == 1
    assert stats["list_items"] == 5
    assert stats["rules"] == 1

    md = document_to_markdown(path)
    assert "# Quarterly Report" in md
    assert "## Summary" in md
    assert "**revenue**" in md
    assert "*growth*" in md
    assert "[reference](https://example.com/data)" in md
    assert "- First finding" in md
    assert "| North | 100 |" in md
    assert "Closing paragraph." in md


def test_created_document_validates_clean(tmp_path: Path):
    path = str(tmp_path / "clean.docx")
    markdown_to_document(SAMPLE_MARKDOWN, path)
    report = validate_docx(path)
    assert report["valid"], report
    assert report["errors"] == []


def test_validation_flags_skewed_table(tmp_path: Path):
    path = str(tmp_path / "skewed.docx")
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(2)
    # Deliberately set the first cell's width far from its grid column.
    table.rows[0].cells[0].width = Inches(4)
    doc.save(path)

    report = validate_docx(path)
    assert any("will skew" in w for w in report["warnings"]), report


def test_validation_rejects_non_docx(tmp_path: Path):
    path = tmp_path / "not_a_doc.docx"
    path.write_text("plain text")
    report = validate_docx(str(path))
    assert not report["valid"]
    assert any("PACKAGE" in e for e in report["errors"])


def test_revision_markup_in_markdown(tmp_path: Path):
    from word_mcp_live_gavintomlins.tools.tracked_changes_tools import track_replace

    path = str(tmp_path / "tracked.docx")
    markdown_to_document("Original wording stays here.", path)
    result = asyncio.run(track_replace(path, "Original wording", "Revised wording"))
    assert "error" not in str(result).lower(), result

    md = document_to_markdown(path, show_revisions=True)
    assert "{++Revised wording++}" in md
    assert "{--Original wording--}" in md

    final = document_to_markdown(path, show_revisions=False)
    assert "Revised wording" in final
    assert "Original wording" not in final


def test_set_update_fields_on_open(tmp_path: Path):
    path = str(tmp_path / "fields.docx")
    markdown_to_document("# Doc\n\nBody.", path)
    message = asyncio.run(quality_tools.set_update_fields_on_open(path))
    assert "refresh fields" in message
    with zipfile.ZipFile(path) as zf:
        assert b"updateFields" in zf.read("word/settings.xml")


def test_validate_document_tool_returns_json(tmp_path: Path):
    path = str(tmp_path / "tool.docx")
    markdown_to_document("# Title\n\nText.", path)
    raw = asyncio.run(quality_tools.validate_document(path))
    report = json.loads(raw)
    assert report["valid"] is True


def _make_template(path: Path, strip_styles=()) -> None:
    """Create a template with a marker footer and a customised Heading 1."""
    doc = Document()
    doc.styles["Heading 1"].font.name = "Courier New"
    doc.sections[0].footer.paragraphs[0].text = "TPL-FOOTER"
    doc.add_paragraph("template body content that must not survive")
    for name in strip_styles:
        doc.styles[name].delete()
    doc.save(str(path))


def test_create_from_template_inherits_styles(tmp_path: Path):
    tpl = tmp_path / "brand_template.docx"
    _make_template(tpl)
    out = str(tmp_path / "branded.docx")
    markdown_to_document(SAMPLE_MARKDOWN, out, template=str(tpl))

    doc = Document(out)
    assert doc.styles["Heading 1"].font.name == "Courier New"
    assert "TPL-FOOTER" in doc.sections[0].footer.paragraphs[0].text
    md = document_to_markdown(out)
    assert "# Quarterly Report" in md
    assert "template body content" not in md
    assert validate_docx(out)["valid"]


def test_create_from_template_without_list_or_table_styles(tmp_path: Path):
    tpl = tmp_path / "sparse_template.docx"
    _make_template(
        tpl,
        strip_styles=("List Bullet", "List Bullet 2", "List Number", "Table Grid"),
    )
    out = str(tmp_path / "sparse.docx")
    stats = markdown_to_document(SAMPLE_MARKDOWN, out, template=str(tpl))
    assert stats["list_items"] == 5 and stats["tables"] == 1

    md = document_to_markdown(out)
    assert "First finding" in md and "Step one" in md
    assert "| North | 100 |" in md
    assert validate_docx(out)["valid"]


def test_create_from_markdown_tool_rejects_missing_template(tmp_path: Path):
    result = asyncio.run(quality_tools.create_document_from_markdown(
        str(tmp_path / "x.docx"), "# T", template=str(tmp_path / "nope.docx")))
    assert "does not exist" in result
